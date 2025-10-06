"""
Document Reader Utilities
Extract text content from various document formats for AI processing.
"""

import os
from typing import Dict, Optional, List
import logging

# PDF reading
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word document reading
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Excel reading
try:
    from openpyxl import load_workbook
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# Image text extraction (OCR would require pytesseract)
try:
    from PIL import Image
    IMAGE_AVAILABLE = True
except ImportError:
    IMAGE_AVAILABLE = False

# Web scraping
try:
    from bs4 import BeautifulSoup
    import requests
    WEB_AVAILABLE = True
except ImportError:
    WEB_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentReader:
    """Read and extract text from various document formats."""
    
    @staticmethod
    def read_pdf(file_path: str, max_pages: int = 50) -> Dict[str, any]:
        """Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            max_pages: Maximum number of pages to read (to prevent memory issues)
            
        Returns:
            Dictionary with text content and metadata
        """
        if not PDF_AVAILABLE:
            return {"error": "PyPDF2 not installed", "text": ""}
        
        try:
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            pages_to_read = min(total_pages, max_pages)
            
            text_content = []
            for i in range(pages_to_read):
                page = reader.pages[i]
                text_content.append(f"\n--- Page {i+1} ---\n{page.extract_text()}")
            
            full_text = "\n".join(text_content)
            
            return {
                "text": full_text,
                "pages": total_pages,
                "pages_read": pages_to_read,
                "truncated": total_pages > max_pages,
                "metadata": reader.metadata if hasattr(reader, 'metadata') else {},
                "success": True
            }
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {str(e)}")
            return {
                "error": str(e),
                "text": "",
                "success": False
            }
    
    @staticmethod
    def read_docx(file_path: str) -> Dict[str, any]:
        """Extract text from Word document.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with text content and metadata
        """
        if not DOCX_AVAILABLE:
            return {"error": "python-docx not installed", "text": ""}
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                tables_text.append("\n".join(table_data))
            
            full_text = "\n\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n=== Tables ===\n\n" + "\n\n".join(tables_text)
            
            return {
                "text": full_text,
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
                "success": True
            }
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {str(e)}")
            return {
                "error": str(e),
                "text": "",
                "success": False
            }
    
    @staticmethod
    def read_excel(file_path: str, max_rows: int = 500) -> Dict[str, any]:
        """Extract data from Excel file.
        
        Args:
            file_path: Path to Excel file
            max_rows: Maximum rows to read per sheet
            
        Returns:
            Dictionary with text representation and metadata
        """
        if not EXCEL_AVAILABLE:
            return {"error": "openpyxl not installed", "text": ""}
        
        try:
            workbook = load_workbook(file_path, read_only=True, data_only=True)
            
            sheets_text = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                sheet_text = [f"\n=== Sheet: {sheet_name} ===\n"]
                
                # Read up to max_rows
                row_count = 0
                for row in sheet.iter_rows(values_only=True):
                    if row_count >= max_rows:
                        sheet_text.append(f"\n... (truncated at {max_rows} rows)")
                        break
                    
                    # Convert row to string, filtering None values
                    row_str = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_str.strip():
                        sheet_text.append(row_str)
                    row_count += 1
                
                sheets_text.append("\n".join(sheet_text))
            
            full_text = "\n\n".join(sheets_text)
            
            return {
                "text": full_text,
                "sheets": len(workbook.sheetnames),
                "sheet_names": workbook.sheetnames,
                "success": True
            }
        except Exception as e:
            logger.error(f"Error reading Excel {file_path}: {str(e)}")
            return {
                "error": str(e),
                "text": "",
                "success": False
            }
    
    @staticmethod
    def read_text_file(file_path: str) -> Dict[str, any]:
        """Read plain text file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            Dictionary with text content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            return {
                "text": text,
                "lines": len(text.split('\n')),
                "success": True
            }
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                return {
                    "text": text,
                    "lines": len(text.split('\n')),
                    "encoding": "latin-1",
                    "success": True
                }
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {str(e)}")
                return {
                    "error": str(e),
                    "text": "",
                    "success": False
                }
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {str(e)}")
            return {
                "error": str(e),
                "text": "",
                "success": False
            }
    
    @staticmethod
    def read_image(file_path: str) -> Dict[str, any]:
        """Get image metadata (OCR would require pytesseract).
        
        Args:
            file_path: Path to image file
            
        Returns:
            Dictionary with image metadata
        """
        if not IMAGE_AVAILABLE:
            return {"error": "Pillow not installed", "text": ""}
        
        try:
            with Image.open(file_path) as img:
                return {
                    "text": f"[Image: {img.format} format, {img.size[0]}x{img.size[1]} pixels, {img.mode} mode]",
                    "format": img.format,
                    "size": img.size,
                    "mode": img.mode,
                    "info": "OCR not implemented - would require pytesseract",
                    "success": True
                }
        except Exception as e:
            logger.error(f"Error reading image {file_path}: {str(e)}")
            return {
                "error": str(e),
                "text": "",
                "success": False
            }
    
    @staticmethod
    def read_url(url: str, timeout: int = 10) -> Dict[str, any]:
        """Fetch and extract text from URL.
        
        Args:
            url: URL to fetch
            timeout: Request timeout in seconds
            
        Returns:
            Dictionary with extracted text and metadata
        """
        if not WEB_AVAILABLE:
            return {"error": "BeautifulSoup not installed", "text": ""}
        
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(url, timeout=timeout, headers=headers)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'lxml')
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            # Limit size
            if len(text) > 10000:
                text = text[:10000] + "\n\n... (truncated at 10,000 characters)"
            
            return {
                "text": text,
                "url": url,
                "title": soup.title.string if soup.title else "No title",
                "status_code": response.status_code,
                "truncated": len(response.text) > 10000,
                "success": True
            }
        except requests.Timeout:
            return {
                "error": "Request timed out",
                "text": "",
                "success": False
            }
        except requests.RequestException as e:
            logger.error(f"Error fetching URL {url}: {str(e)}")
            return {
                "error": str(e),
                "text": "",
                "success": False
            }
        except Exception as e:
            logger.error(f"Error parsing URL {url}: {str(e)}")
            return {
                "error": str(e),
                "text": "",
                "success": False
            }
    
    @staticmethod
    def read_document(file_path: str) -> Dict[str, any]:
        """Auto-detect file type and extract content.
        
        Args:
            file_path: Path to document
            
        Returns:
            Dictionary with extracted content
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return DocumentReader.read_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return DocumentReader.read_docx(file_path)
        elif ext in ['.xlsx', '.xls']:
            return DocumentReader.read_excel(file_path)
        elif ext in ['.txt', '.csv', '.md', '.log']:
            return DocumentReader.read_text_file(file_path)
        elif ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
            return DocumentReader.read_image(file_path)
        else:
            return {
                "error": f"Unsupported file type: {ext}",
                "text": "",
                "success": False
            }


def get_attachment_content(file_metadata: Dict, url_metadata: List[Dict]) -> str:
    """
    Get combined content from file attachments and URLs for AI context.
    
    Args:
        file_metadata: File attachment metadata
        url_metadata: URL reference metadata
        
    Returns:
        Combined text content string
    """
    content_parts = []
    
    # Process files
    if file_metadata:
        content_parts.append("=== Attached Documents ===\n")
        for file_meta in file_metadata:
            file_path = file_meta.get('file_path')
            filename = file_meta.get('filename', 'Unknown')
            
            if file_path and os.path.exists(file_path):
                result = DocumentReader.read_document(file_path)
                if result.get('success'):
                    content_parts.append(f"\n--- Document: {filename} ---\n{result['text']}\n")
                else:
                    content_parts.append(f"\n--- Document: {filename} ---\n[Error reading: {result.get('error', 'Unknown error')}]\n")
    
    # Process URLs
    if url_metadata:
        content_parts.append("\n=== Referenced URLs ===\n")
        for url_meta in url_metadata:
            url = url_meta.get('url')
            title = url_meta.get('title', 'Unknown')
            
            if url:
                result = DocumentReader.read_url(url)
                if result.get('success'):
                    content_parts.append(f"\n--- URL: {title} ({url}) ---\n{result['text']}\n")
                else:
                    content_parts.append(f"\n--- URL: {title} ({url}) ---\n[Error fetching: {result.get('error', 'Unknown error')}]\n")
    
    return "\n".join(content_parts) if content_parts else ""
